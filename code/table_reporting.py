#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Table Reporting Module

This module generates tables and figures from processed timber asset data.
It follows the single responsibility principle by focusing only on visualization
and reporting, with data processing handled by region-specific modules.

The module can generate reports for different regions:
- Southern region
- Great Lakes region
- Combined national reports
"""

import pandas as pd
import matplotlib.pyplot as plt
import docx
from pathlib import Path
from geo_crosswalks import DATA_DIR
from species_crosswalks import convert_to_billions, convert_to_megatonnes


def load_processed_data(region):
    """
    Load processed data files for a specific region.
    
    Parameters:
    -----------
    region : str
        Region identifier ('south', 'gl', or 'all')
        
    Returns:
    --------
    pandas.DataFrame
        Processed table data
    """
    if region.lower() == 'south':
        return pd.read_csv(DATA_DIR / 'processed' / 'table_south.csv')
    elif region.lower() in ('gl', 'great_lakes'):
        return pd.read_csv(DATA_DIR / 'processed' / 'table_gl.csv')
    elif region.lower() == 'all':
        # Load and combine both regions
        south_data = pd.read_csv(DATA_DIR / 'processed' / 'table_south.csv')
        gl_data = pd.read_csv(DATA_DIR / 'processed' / 'table_gl.csv')
        
        # Add region identifier
        south_data['region'] = 'South'
        gl_data['region'] = 'Great Lakes'
        
        # Combine
        return pd.concat([south_data, gl_data], ignore_index=True)
    else:
        raise ValueError(f"Unknown region: {region}. Must be 'south', 'gl', or 'all'.")


def add_species_info(table_data):
    """
    Add species and species group information to the table data.
    
    Parameters:
    -----------
    table_data : pandas.DataFrame
        Processed table data
        
    Returns:
    --------
    pandas.DataFrame
        Table data with added species information
    """
    # Load species and species group dictionaries
    species = pd.read_csv(DATA_DIR / 'speciesDict.csv')
    species_group = pd.read_csv(DATA_DIR / 'speciesGroupDict.csv')
    
    # Merge with species and species group information
    table_data = pd.merge(table_data, species, on='spcd')
    table_data = pd.merge(table_data, species_group, on='spgrpcd')
    
    # Standardize species class names
    table_data = table_data.replace({'spclass': {'Softwood': 'Coniferous',
                                              'Hardwood': 'Non-coniferous',
                                              'Pine': 'Coniferous',
                                              'Oak': 'Non-coniferous'}})
    
    return table_data


def generate_physical_table_by_species_class(table_data, region_name):
    """
    Generate physical table (volume) by species class.
    
    Parameters:
    -----------
    table_data : pandas.DataFrame
        Processed table data with species information
    region_name : str
        Name of the region for table title
        
    Returns:
    --------
    pandas.DataFrame
        Formatted table
    """
    # Filter out pre-merchantable if present
    filtered_data = table_data[table_data['product'] != 'Pre-merchantable']
    
    # Group by species class and product
    table = filtered_data.groupby(['spclass', 'product']).agg(
        {'volume': 'sum'}).reset_index()
    
    # Pivot to put products in columns
    table = table.pivot(index='spclass',
                        columns='product',
                        values='volume').reset_index()
    
    # Sort by species class
    table = table.sort_values('spclass', ascending=True)
    
    return table


def generate_monetary_table_by_species_class(table_data, region_name):
    """
    Generate monetary table (value) by species class.
    
    Parameters:
    -----------
    table_data : pandas.DataFrame
        Processed table data with species information
    region_name : str
        Name of the region for table title
        
    Returns:
    --------
    pandas.DataFrame
        Formatted table
    """
    # Filter out pre-merchantable if present
    filtered_data = table_data[table_data['product'] != 'Pre-merchantable']
    
    # Group by species class and product
    table = filtered_data.groupby(['spclass', 'product']).agg(
        {'value': 'sum'}).reset_index()
    
    # Pivot to put products in columns
    table = table.pivot(index='spclass',
                        columns='product',
                        values='value').reset_index()
    
    # Sort by species class
    table = table.sort_values('spclass', ascending=True)
    
    return table


def generate_physical_table_by_species_group(table_data, region_name):
    """
    Generate physical table (volume) by species group.
    
    Parameters:
    -----------
    table_data : pandas.DataFrame
        Processed table data with species information
    region_name : str
        Name of the region for table title
        
    Returns:
    --------
    pandas.DataFrame
        Formatted table
    """
    # Filter out pre-merchantable if present
    filtered_data = table_data[table_data['product'] != 'Pre-merchantable']
    
    # Group by species group and product
    table = filtered_data.groupby(['speciesGroup', 'product']).agg(
        {'volume': 'sum'}).reset_index()
    
    # Pivot to put products in columns
    table = table.pivot(index='speciesGroup',
                        columns='product',
                        values='volume').reset_index()
    
    # Sort by sawtimber volume in descending order
    if 'Sawtimber' in table.columns:
        table = table.sort_values('Sawtimber', ascending=False)
    
    return table


def generate_monetary_table_by_species_group(table_data, region_name):
    """
    Generate monetary table (value) by species group and class.
    
    Parameters:
    -----------
    table_data : pandas.DataFrame
        Processed table data with species information
    region_name : str
        Name of the region for table title
        
    Returns:
    --------
    pandas.DataFrame
        Formatted table
    """
    # Filter out pre-merchantable if present
    filtered_data = table_data[table_data['product'] != 'Pre-merchantable']
    
    # Group by species class, species group, and product
    table = filtered_data.groupby(['spclass', 'speciesGroup', 'product']).agg(
        {'value': 'sum'}).reset_index()
    
    # Sort by value in descending order
    table = table.sort_values(by='value', ascending=False)
    
    return table


def generate_tables_for_region(region):
    """
    Generate all tables for a specific region.
    
    Parameters:
    -----------
    region : str
        Region identifier ('south', 'gl', or 'all')
        
    Returns:
    --------
    dict
        Dictionary containing all generated tables
    """
    # Set region name for titles
    if region.lower() == 'south':
        region_name = 'South'
    elif region.lower() in ('gl', 'great_lakes'):
        region_name = 'Great Lakes'
    else:
        region_name = 'US'
    
    # Load data
    table_data = load_processed_data(region)
    
    # Add species information
    table_data = add_species_info(table_data)
    
    # Generate tables
    table1 = generate_physical_table_by_species_class(table_data, region_name)
    table2 = generate_monetary_table_by_species_class(table_data, region_name)
    table3 = generate_physical_table_by_species_group(table_data, region_name)
    table4 = generate_monetary_table_by_species_group(table_data, region_name)
    
    return {
        'physical_by_species_class': table1,
        'monetary_by_species_class': table2,
        'physical_by_species_group': table3,
        'monetary_by_species_group': table4
    }


def create_word_document(tables, region):
    """
    Create a Word document with all tables.
    
    Parameters:
    -----------
    tables : dict
        Dictionary containing all generated tables
    region : str
        Region identifier ('south', 'gl', or 'all')
        
    Returns:
    --------
    None
    """
    # Set region name for titles
    if region.lower() == 'south':
        region_name = 'South'
    elif region.lower() in ('gl', 'great_lakes'):
        region_name = 'Great Lakes'
    else:
        region_name = 'US'
    
    # Create new Word document
    doc = docx.Document()
    doc.add_heading(f'{region_name} Timber Asset Tables', level=1)
    
    # Table list
    table_list = [
        tables['physical_by_species_class'],
        tables['monetary_by_species_class'],
        tables['physical_by_species_group'],
        tables['monetary_by_species_group']
    ]
    
    # Table titles
    table_titles = [
        f'Table 1: {region_name}: Physical Table by Species Class',
        f'Table 2: {region_name}: Monetary Value Table by Species Class',
        f'Table 3: {region_name}: Physical Table by Timber Species Group',
        f'Table 4: {region_name}: Monetary Value Table by Timber Species Group'
    ]
    
    # Function to add a table to the Word document
    def add_table_to_doc(df, title):
        doc.add_heading(title, level=2)
        doc_table = doc.add_table(rows=1, cols=len(df.columns))
        hdr_cells = doc_table.rows[0].cells
        
        for i, col in enumerate(df.columns):
            hdr_cells[i].text = str(col)
        
        for index, row in df.iterrows():
            row_cells = doc_table.add_row().cells
            for i, value in enumerate(row):
                if pd.isna(value):
                    row_cells[i].text = '-'
                elif isinstance(value, (int, float)):
                    if i >= 1:  # Assuming first column is labels
                        if 'Physical' in title:
                            row_cells[i].text = f"{value:,.0f} Mt"
                        else:
                            row_cells[i].text = f"${value:,.2f}B"
                    else:
                        row_cells[i].text = str(value)
                else:
                    row_cells[i].text = str(value)
    
    # Add tables to document
    for df, title in zip(table_list, table_titles):
        add_table_to_doc(df, title)
    
    # Create reports directory if it doesn't exist
    reports_dir = DATA_DIR / 'reports'
    reports_dir.mkdir(exist_ok=True)
    
    # Save document
    doc.save(reports_dir / f'{region_name.lower()}_timber_asset_tables.docx')
    print(f"Word document saved: {region_name.lower()}_timber_asset_tables.docx")


def generate_region_summary_chart(region):
    """
    Generate a summary chart for the region.
    
    Parameters:
    -----------
    region : str
        Region identifier ('south', 'gl', or 'all')
        
    Returns:
    --------
    None
    """
    # Set region name for titles
    if region.lower() == 'south':
        region_name = 'South'
    elif region.lower() in ('gl', 'great_lakes'):
        region_name = 'Great Lakes'
    else:
        region_name = 'US'
    
    # Load data
    table_data = load_processed_data(region)
    
    # Add species information
    table_data = add_species_info(table_data)
    
    # Filter and summarize data
    summary = table_data.groupby(['spclass', 'product']).agg({
        'volume': 'sum',
        'value': 'sum'
    }).reset_index()
    
    # Create figure with two subplots
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(15, 7))
    
    # Plot volume by species class and product
    volume_pivot = summary.pivot(index='spclass', columns='product', values='volume')
    volume_pivot.plot(kind='bar', stacked=True, ax=ax1, 
                      colormap='viridis')
    ax1.set_title(f'{region_name} Timber Volume by Species Class and Product')
    ax1.set_ylabel('Volume (Megatonnes)')
    ax1.grid(axis='y', linestyle='--', alpha=0.7)
    
    # Plot value by species class and product
    value_pivot = summary.pivot(index='spclass', columns='product', values='value')
    value_pivot.plot(kind='bar', stacked=True, ax=ax2,
                     colormap='viridis')
    ax2.set_title(f'{region_name} Timber Value by Species Class and Product')
    ax2.set_ylabel('Value (Billions $)')
    ax2.grid(axis='y', linestyle='--', alpha=0.7)
    
    plt.tight_layout()
    
    # Create reports directory if it doesn't exist
    reports_dir = DATA_DIR / 'reports'
    reports_dir.mkdir(exist_ok=True)
    
    # Save figure
    plt.savefig(reports_dir / f'{region_name.lower()}_timber_summary.png', dpi=300, bbox_inches='tight')
    print(f"Summary chart saved: {region_name.lower()}_timber_summary.png")
    
    plt.close()


def generate_reports(regions=None):
    """
    Generate all reports for specified regions.
    
    Parameters:
    -----------
    regions : list, optional
        List of region identifiers to generate reports for.
        If None, generates reports for all regions.
        
    Returns:
    --------
    None
    """
    if regions is None:
        regions = ['south', 'gl', 'all']
    
    # Create reports directory
    reports_dir = DATA_DIR / 'reports'
    reports_dir.mkdir(exist_ok=True)
    
    for region in regions:
        try:
            print(f"Generating reports for {region}...")
            
            # Generate tables
            tables = generate_tables_for_region(region)
            
            # Create Word document
            create_word_document(tables, region)
            
            # Generate summary chart
            generate_region_summary_chart(region)
            
            print(f"Reports for {region} complete!")
        except Exception as e:
            print(f"Error generating reports for {region}: {e}")


if __name__ == "__main__":
    # Generate reports for all regions
    generate_reports() 
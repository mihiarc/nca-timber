#!/usr/bin/env python
# coding: utf-8

# In[3]:


from utils import speciesDict, get_species_name, speciesGroupDict, get_species_group_name
import pandas as pd
import docx
import matplotlib.pyplot as plt

# This script should only handle table/figure generation.
# Data wrangling and region-specific logic should be handled elsewhere and passed in as arguments or via pre-processed files.

# Example: Load pre-processed southern data for table/figure generation
# (Assume data is already cleaned and merged as needed)
tableSouth = pd.read_csv('../data/tableSouthMerch.csv')
species = pd.read_csv('../data/speciesDict.csv')
speciesGroup = pd.read_csv('../data/speciesGroupDict.csv')

tableSouth = pd.merge(tableSouth, species, on='spcd')
tableSouth = pd.merge(tableSouth, speciesGroup, on='spgrpcd')
tableSouth = tableSouth.replace({'spclass': {'Softwood': 'Coniferous',
                                              'Hardwood': 'Non-coniferous'}})

# scale to billions of dollars
tableSouth['value'] = tableSouth['value'] / 1e9
# convert to cubic feet to megatonnes
tableSouth['volume'] = tableSouth['volume'] * 0.025713 / 1e6


# In[7]:


# create a new Word document
doc = docx.Document()
doc.add_heading('US Timber Asset Pilots', level=1)

# define the dataframes
southTable1 = tableSouth[tableSouth['product'] != 'Pre-merchantable']
southTable1 = southTable1.groupby(['spclass', 'product']).agg(
                    {'volume': 'sum'}).reset_index()
southTable1 = southTable1.pivot(index='spclass',
                                columns='product',
                                values='volume').reset_index()
southTable1 = southTable1.sort_values('spclass', ascending=True)

# reduce value and volume across space
southTable2 = tableSouth[tableSouth['product'] != 'Pre-merchantable']
southTable2 = southTable2.groupby(['spclass', 'product']).agg(
                    {'value': 'sum'}).reset_index()
southTable2 = southTable2.pivot(index='spclass',
                                columns='product',
                                values='value').reset_index()
southTable2 = southTable2.sort_values('spclass', ascending=True)

# reduce value and volume across space
southTable3 = tableSouth[tableSouth['product'] != 'Pre-merchantable']
southTable3 = southTable3.groupby(['speciesGroup', 'product']).agg(
                    {'volume': 'sum'}).reset_index()
southTable3 = southTable3.pivot(index='speciesGroup',
                                columns='product',
                                values='volume').reset_index()
southTable3 = southTable3.sort_values('Sawtimber', ascending=False)

# reduce value and volume across space
southTable4 = tableSouth[tableSouth['product'] != 'Pre-merchantable']
southTable4 = southTable4.groupby(['spclass', 'speciesGroup', 'product']).agg(
                    {'value': 'sum'}).reset_index()
southTable4 = southTable4.sort_values(by='value', ascending=False)

# list the dataframes
tables = [southTable1, southTable2, southTable3, southTable4]

# titles for the tables
titles = [
    'Table 1: South: Physical Table by Species Class',
    'Table 2: South: Monetary Value Table by Species Class',
    'Table 3: South: Physical Table by Timber Species Group',
    'Table 4: South: Monetary Value Table by Timber Species Group'
]

# function to add a table to the Word document
def add_table_to_doc(df, title):
    doc.add_heading(title, level=2)
    doc_table = doc.add_table(rows=1, cols=len(df.columns))
    hdr_cells = doc_table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = str(col)

    for index, row in df.iterrows():
        row_cells = doc_table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = '-' if pd.isna(value) else str(value)

# add the tables to the Word document
for df, title in zip(tables, titles):
    add_table_to_doc(df, title)

# save the Word document
doc.save('../data/tables_and_figures.docx')
